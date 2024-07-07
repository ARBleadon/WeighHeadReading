import os
import time
import csv
from datetime import datetime
import threading
import random
import msvcrt  # For Windows getch() function
import gc
import re
import win32print
import serial  # PySerial library for serial communication

exit_flag = False

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# CSV file path
csv_file_path = os.path.join(script_dir, 'weighing_data.csv')

# Create the CSV file if it doesn't exist and write headers
if not os.path.exists(csv_file_path):
    with open(csv_file_path, 'w', newline='') as csvfile:
        csv_writer = csv.writer(csvfile)
        csv_writer.writerow(['BagID', 'GrossWeight', 'DateandTime', 'BatchNumb', 'ProductType'])

# Extracts the number from a string
def extract_number(input_string):
    pattern = r"\d+(\.\d+)?"
    match = re.search(pattern, input_string)
    if match:
        extracted_number = match.group(0)
        return float(extracted_number)  # Convert to float
    else:
        print("No numbers found in the input string.")
        return None

# Get the last batch number in the CSV
def get_last_batch_number():
    with open(csv_file_path, 'r') as csvfile:
        csv_reader = csv.reader(csvfile)
        next(csv_reader)  # Skip header row
        rows = list(csv_reader)
        if not rows:
            return 1
        last_row = rows[-1]
        last_batch_number = int(last_row[3])
        entries_in_last_batch = sum(1 for row in rows if row[3] == str(last_batch_number))
        if entries_in_last_batch >= 22:
            return last_batch_number + 1
        return last_batch_number

# Count entries in a specific batch
def count_entries_in_batch(batch_number):
    with open(csv_file_path, 'r') as csvfile:
        csv_reader = csv.reader(csvfile)
        next(csv_reader)  # Skip header row
        return sum(1 for row in csv_reader if row[3] == str(batch_number))

# Print label to printer
def print_file_to_printer(weight, bag_id):
    try:
        printer_name = "ZDesigner GK420d"  # Adjust to your printer name
        zpl = f"""
        ^XA
        ^PW800
        ^LL600
        ^FO100,100
        ^A0N,100,100
        ^FDWeight: {weight}^FS
        ^FO100,250
        ^A0N,100,100
        ^FDBagID: {bag_id}^FS
        ^XZ
        """
        
        hPrinter = win32print.OpenPrinter(printer_name)
        try:
            hJob = win32print.StartDocPrinter(hPrinter, 1, ("Label", None, "RAW"))
            try:
                win32print.StartPagePrinter(hPrinter)
                win32print.WritePrinter(hPrinter, zpl.encode('utf-8'))
                win32print.EndPagePrinter(hPrinter)
            finally:
                win32print.EndDocPrinter(hPrinter)
        finally:
            win32print.ClosePrinter(hPrinter)
            
        print("Print job sent successfully")
    except Exception as e:
        print(f"An error occurred: Unable to open printer ({e})")

# Print report to Word document
def print_file_to_word_doc(file_path, report_date):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Apply table style

    # Add header row with column names
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'BagID'
    hdr_cells[1].text = 'GrossWeight'
    hdr_cells[2].text = 'DateandTime'
    hdr_cells[3].text = 'BatchNumb'
    hdr_cells[4].text = 'ProductType'

    try:
        with open(file_path, 'r') as file:
            csv_reader = csv.reader(file)
            next(csv_reader)  # Skip header row
            for row in csv_reader:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = cell_value
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        return
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return

    # Format table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

    # Center align text in all cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Save Word document with date and time in the script directory
    doc_filename = f'report_{report_date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
    doc.save(os.path.join(script_dir, doc_filename))  
    print(f"Word document saved to {os.path.join(script_dir, doc_filename)}")

# Read data from serial port
def read_serial_data(port, baud_rate, ser=None):
    print(f"Reading data from serial port {port}... Press '0' to return to the main menu.")

    if ser is None:
        ser = serial.Serial(port, baud_rate, timeout=1)  # Open serial port
        print(f"Opened serial port {port} successfully.")

    try:
        read_count = 0  
        batch_number = get_last_batch_number()
        entries_in_batch = count_entries_in_batch(batch_number)

        while entries_in_batch < 22 and not exit_flag:
            data = ser.readline().decode('utf-8').strip()  # Read serial data

            if data.startswith("Gross"):
                number = extract_number(data)
                read_count += 1

                current_time = datetime.now()

                # Generate random BagID
                random_bag_id = ''.join(random.choices('0123456789', k=6))

                with open(csv_file_path, 'a', newline='') as csvfile:
                    csv_writer = csv.writer(csvfile)
                    csv_writer.writerow([random_bag_id, number, current_time, batch_number, 'Product'])

                print(f"\n\nBatch: {batch_number}   Weight: {number}    BagID: {random_bag_id}       Date and time: {current_time}")
                print(f"Entries in current batch: {entries_in_batch + 1}\n")

                print_file_to_printer(number, random_bag_id)
                entries_in_batch += 1

                if entries_in_batch == 22:
                    batch_number += 1
                    entries_in_batch = 0

            elif data == '0':
                print("Returned to the main menu.")
                break

    except KeyboardInterrupt:
        print("Interrupted by user. Exiting to menu.")

    except Exception as e:
        print(f"An error occurred: {e}")

    finally:
        if ser:
            ser.close()  # Close serial port
            time.sleep(2)  # Wait for 2 seconds before re-opening
            print(f"Closed serial port {port}")

            # Reopen menu
            menu(port, baud_rate)

# Update bag weight
def update(port, baud_rate, ser):
    print("Update is used if a bag is broken or incorrectly weighed\n")
    print("Please Enter the Bag ID for the bag that needs reweighing\n")

    while True:
        bag_id = input("-->")
        if bag_id.isdigit():
            break
        else:
            print("Invalid input. Please enter a numeric Bag ID.")

    ser = serial.Serial(port, baud_rate, timeout=1)  # Open serial port
    print(f"Opened serial port {port} successfully.")

    try:
        while True:
            data = ser.readline().decode('utf-8').strip()  # Read serial data

            if data.startswith("Gross"):
                new_weight = extract_number(data)
                current_time = datetime.now()

                with open(csv_file_path, 'r') as csvfile:
                    csv_reader = csv.reader(csvfile)
                    rows = list(csv_reader)

                for row in rows:
                    if row[0] == bag_id:
                        row[1] = new_weight
                        row[2] = current_time
                        break

                with open(csv_file_path, 'w', newline='') as csvfile:
                    csv_writer = csv.writer(csvfile)
                    csv_writer.writerows(rows)

                print(f"Bag ID {bag_id} reweighed successfully: New weight {new_weight} kg at {current_time}")

                return

    except KeyboardInterrupt:
        print("Interrupted by user. Exiting to menu.")

    except Exception as e:
        print(f"An error occurred: {e}")

    finally:
        if ser:
            ser.close()  # Close serial port
            time.sleep(2)  # Wait for 2 seconds before re-opening
            print(f"Closed serial port {port}")
            menu(port, baud_rate)

# View bag details
def view_bag_details():
    while True:
        bag_id = input("Enter the Bag ID to view details: ")
        if bag_id.isdigit():
            break
        else:
            print("Invalid input. Please enter a numeric Bag ID.")

    with open(csv_file_path, 'r') as csvfile:
        csv_reader = csv.reader(csvfile)
        next(csv_reader)  # Skip header row
        found = False
        for row in csv_reader:
            if row[0] == bag_id:
                print(f"Details for BagID {bag_id}:")
                print(f"BagID: {row[0]}")
                print(f"GrossWeight: {row[1]}")
                print(f"Date and Time: {row[2]}")
                print(f"Batch Number: {row[3]}")
                print(f"Product Type: {row[4]}")
                found = True
                break

        if not found:
            print(f"No details found for BagID {bag_id}")

# Menu function
def menu(port, baud_rate):
    global exit_flag
    exit_flag = False

    while True:
        print("Choose an option:")
        print("1) Read Serial Data")
        print("2) Update Bag Weight")
        print("3) View Bag Details")
        print("4) Print Report")
        print("5) Exit")
        choice = input("Enter your choice: ")

        if choice == '1':
            read_serial_data(port, baud_rate)
        elif choice == '2':
            update(port, baud_rate, None)
        elif choice == '3':
            view_bag_details()
        elif choice == '4':
            report_date = datetime.now()
            print_file_to_word_doc(csv_file_path, report_date)
        elif choice == '5':
            print("Exiting...")
            exit_flag = True  # Signal to stop serial data simulation
            break
        else:
            print("Invalid choice. Please enter a number between 1 and 5.")

if __name__ == "__main__":
    port = 'COM3'  # Placeholder for the actual serial port
    baud_rate = 9600
    menu(port, baud_rate)
